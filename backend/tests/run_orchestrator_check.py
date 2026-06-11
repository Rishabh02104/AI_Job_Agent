import os
import sys
import docx

# Add backend directory to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from utils.parser import extract_text_from_file, get_embedding, parse_resume_json
from orchestrator import Orchestrator
from db import supabase

def create_sample_resume_docx(file_path: str):
    """
    Generates a clean resume matching the tech stack.
    """
    doc = docx.Document()
    doc.add_heading("Rishabh Sharma", level=0)
    doc.add_paragraph("AI Software Engineer | risha@example.com | Bengaluru, India")
    
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Programming Languages: Python, JavaScript, TypeScript, SQL, HTML/CSS")
    doc.add_paragraph("Frameworks & Tools: FastAPI, Next.js, TailwindCSS, React, Supabase, pgvector, Node.js")
    doc.add_paragraph("AI & Automation: Groq API, Llama models, sentence-transformers, Playwright, reportlab")
    
    doc.add_heading("Professional Experience", level=1)
    p1 = doc.add_paragraph()
    p1.add_run("AI Engineer at CareerForge ").bold = True
    p1.add_run("(2024 - Present)\n").italic = True
    p1.add_run("- Built and deployed multi-agent platforms optimizing resumes and cover letters.\n")
    p1.add_run("- Developed FastAPI backend pipelines integrated with Groq Llama inference models.\n")
    p1.add_run("- Configured Supabase PostgreSQL storage with pgvector for semantic search.")
    
    doc.save(file_path)

def run_orchestrator_test():
    print("="*60)
    print("AI Job Agent - E2E Orchestrator Pipeline Verification")
    print("="*60)

    # 1. Clean existing records in test environment
    print("\n[Step 0] Cleaning database test records...")
    try:
        supabase.table("matches").delete().neq("id", "00000000-0000-0000-0000-000000000000").execute()
        supabase.table("applications").delete().neq("id", "00000000-0000-0000-0000-000000000000").execute()
        supabase.table("jobs").delete().neq("id", "00000000-0000-0000-0000-000000000000").execute()
        supabase.table("resumes").delete().neq("id", "00000000-0000-0000-0000-000000000000").execute()
        print("Database cleaned successfully.")
    except Exception as e:
        print(f"Warning: database cleaning encountered errors: {e}")

    # 2. Upload Candidate Resume
    resume_path = "sample_resume.docx"
    create_sample_resume_docx(resume_path)
    
    print("\n[Step 1] Loading and embedding candidate resume...")
    raw_text = extract_text_from_file(resume_path)
    parsed_json = parse_resume_json(raw_text)
    resume_embedding = get_embedding(raw_text)
    
    supabase.table("resumes").insert({
        "raw_text": raw_text,
        "parsed_json": parsed_json,
        "embedding": resume_embedding
    }).execute()
    print("Resume loaded and saved to Supabase.")
    
    if os.path.exists(resume_path):
        os.remove(resume_path)

    # 3. Create a PERFECTLY matching job description to force match score > 80%
    print("\n[Step 2] Injecting a highly-matching mock job description...")
    job_desc = (
        "We are looking for an AI Software Engineer who is an expert in Python, FastAPI, and Next.js. "
        "The ideal candidate must have experience building agentic workflows with Groq API and Llama models, "
        "managing database schemas using Supabase and PostgreSQL pgvector, and generating automated PDF resumes "
        "using reportlab. Skills in TypeScript and Playwright automation are required."
    )
    job_embedding = get_embedding(job_desc)
    job_res = supabase.table("jobs").insert({
        "title": "Senior AI Software Engineer (FastAPI/Next.js)",
        "company": "NextGen AI Labs",
        "description": job_desc,
        "location": "Remote, India",
        "source": "Mock Injector",
        "url": "https://nextgenailabs.example.com/jobs/senior-ai-engineer-1",
        "embedding": job_embedding
    }).execute()
    job_id = job_res.data[0]["id"]
    print(f"Mock matching job injected with ID: {job_id}")

    # 4. Trigger Orchestrator
    print("\n[Step 3] Running Orchestrator pipeline...")
    orchestrator = Orchestrator()
    # Run pipeline with a search keyword (Scout will search, Scorer will run on all, and Tailoring will run on the queued job)
    status = orchestrator.run_pipeline(keywords="AI Engineer", limit=1)
    
    print("\n[Step 4] Pipeline completion summary:")
    print(f"Scout fetched: {status['scout']}")
    print(f"Scorer details: {status['scorer']}")
    print(f"Tailored application IDs: {status['tailored_applications']}")
    print(f"Pipeline errors: {status['errors']}")

    # 5. Check if the application advanced to 'reviewing' and contains the tailored assets
    print("\n[Step 5] Checking staged review package details in Supabase...")
    app_res = supabase.table("applications").select("*, jobs(title, company)").eq("job_id", job_id).execute()
    if app_res.data:
        app = app_res.data[0]
        print(f"Staged Application status: {app.get('status')} (Expected: reviewing)")
        print(f"Tailored Resume PDF URL: {app.get('tailored_resume_url')}")
        print("\nTailored Cover Letter Draft:")
        print("-" * 50)
        print(app.get("cover_letter"))
        print("-" * 50)
        
        # Verify local PDF file exists
        local_pdf_path = app.get('tailored_resume_url', '').replace("file:///", "").replace("/", os.sep)
        if os.path.exists(local_pdf_path):
            print(f"Success: Tailored PDF resume generated locally at: {local_pdf_path}")
        else:
            print(f"Error: Tailored PDF resume not found at: {local_pdf_path}")
    else:
        print("Error: No application staged for the mock matching job.")

    print("="*60)
    print("Orchestrator Backend Verification Complete!")
    print("="*60)

if __name__ == "__main__":
    run_orchestrator_test()
