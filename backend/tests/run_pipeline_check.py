import os
import sys
import docx

# Add backend directory to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from utils.parser import extract_text_from_file, get_embedding, parse_resume_json
from agents.scout import ScoutAgent
from agents.scorer import ScorerAgent
from db import supabase
from config import settings

def create_sample_resume_docx(file_path: str):
    """
    Helper to generate a clean .docx resume file using python-docx.
    """
    print(f"Creating sample resume document at: {file_path}")
    doc = docx.Document()
    
    doc.add_heading("Rishabh Sharma", level=0)
    doc.add_paragraph("AI Software Engineer | risha@example.com | Bengaluru, India")
    
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Programming Languages: Python, JavaScript, TypeScript, SQL, HTML/CSS")
    doc.add_paragraph("Frameworks & Tools: FastAPI, Next.js, TailwindCSS, React, Supabase, pgvector, Node.js")
    doc.add_paragraph("AI & Automation: Groq API, Llama models, sentence-transformers, Playwright, BeautifulSoup, pdfplumber")
    
    doc.add_heading("Professional Experience", level=1)
    p1 = doc.add_paragraph()
    p1.add_run("AI Engineer at CareerForge ").bold = True
    p1.add_run("(2024 - Present)\n").italic = True
    p1.add_run("- Built and deployed multi-agent platforms optimizing resumes and cover letters.\n")
    p1.add_run("- Developed FastAPI backend pipelines integrated with Groq Llama inference models.\n")
    p1.add_run("- Configured Supabase PostgreSQL storage with pgvector for semantic search and embeddings.")

    doc.add_heading("Education", level=1)
    p2 = doc.add_paragraph()
    p2.add_run("Bachelor of Technology in Computer Science ").bold = True
    p2.add_run("- National Institute of Technology, Trichy (2020 - 2024)").italic = True
    
    doc.add_heading("Projects", level=1)
    p3 = doc.add_paragraph()
    p3.add_run("Autonomous Job Applicator\n").bold = True
    p3.add_run("- Orchestrated browser bots with Playwright to fill job applications and track responses.\n")
    p3.add_run("- Implemented local SentenceTransformers text embeddings to match candidate profiles to job boards.")
    
    doc.save(file_path)

def run_integration_pipeline():
    print("="*60)
    print("AI Job Agent - E2E Backend Pipeline Verification")
    print("="*60)

    # 1. Clean existing records in test environment
    print("\n[Step 0] Cleaning existing database test records...")
    try:
        supabase.table("matches").delete().neq("id", "00000000-0000-0000-0000-000000000000").execute()
        supabase.table("applications").delete().neq("id", "00000000-0000-0000-0000-000000000000").execute()
        supabase.table("jobs").delete().neq("id", "00000000-0000-0000-0000-000000000000").execute()
        supabase.table("resumes").delete().neq("id", "00000000-0000-0000-0000-000000000000").execute()
        print("Database cleaned successfully.")
    except Exception as e:
        print(f"Warning: database cleaning encountered errors (might be due to empty tables): {e}")

    # 2. Generate and parse resume
    resume_path = "sample_resume.docx"
    create_sample_resume_docx(resume_path)
    
    print("\n[Step 1] Reading and parsing resume file...")
    raw_text = extract_text_from_file(resume_path)
    print("Text extracted. Parsing resume content to JSON via Groq...")
    
    parsed_json = parse_resume_json(raw_text)
    print(f"Parsed Resume JSON:\n{parsed_json}")
    
    print("Generating local 384-d embedding vector...")
    resume_embedding = get_embedding(raw_text)
    
    print("Inserting resume record into Supabase...")
    resume_insert_res = supabase.table("resumes").insert({
        "raw_text": raw_text,
        "parsed_json": parsed_json,
        "embedding": resume_embedding
    }).execute()
    
    resume_id = resume_insert_res.data[0]["id"]
    print(f"Resume saved successfully with ID: {resume_id}")
    
    # Clean up local test file
    if os.path.exists(resume_path):
        os.remove(resume_path)

    # 3. Fetch jobs using Job Scout Agent
    print("\n[Step 2] Launching Scout Agent to pull job listings...")
    scout = ScoutAgent()
    scout_result = scout.run({"keywords": "AI Engineer", "limit": 3})
    print(f"Scout completion details: {scout_result.data}")
    
    if not scout_result.success or scout_result.data.get("saved", 0) == 0:
        print("Scout did not save any new jobs. Running fallback mock job inserts to ensure pipeline runs...")
        # Fallback Mock Job to continue verification
        mock_embedding = get_embedding("We are looking for a Software Engineer with Python and FastAPI experience.")
        supabase.table("jobs").insert({
            "title": "Software Developer (Python/FastAPI)",
            "company": "FastAPI Tech Ltd",
            "description": "Looking for a Software Engineer proficient in Python, FastAPI, and Supabase. Experience with AI embeddings is a plus.",
            "location": "Bengaluru, India",
            "source": "Mock Scout",
            "url": "https://example.com/jobs/mock-fastapi-1",
            "embedding": mock_embedding
        }).execute()
        print("Inserted mock job for verification.")

    # 4. Score jobs using Scorer Agent
    print("\n[Step 3] Launching Scorer Agent...")
    scorer = ScorerAgent()
    scorer_result = scorer.run({"limit": 5})
    
    print(f"Scorer Success: {scorer_result.success}")
    if scorer_result.success:
        print(f"Scored Job details:\n{scorer_result.data}")
    else:
        print(f"Scorer failed: {scorer_result.error}")
        
    print("\n[Step 4] Checking staged applications...")
    apps_res = supabase.table("applications").select("*, jobs(title, company)").execute()
    print(f"Staged applications in queue: {len(apps_res.data)}")
    for app in apps_res.data:
        job = app.get("jobs", {})
        print(f"- Job: {job.get('title')} at {job.get('company')} | Status: {app.get('status')}")

    print("="*60)
    print("Backend pipeline E2E verification complete!")
    print("="*60)

if __name__ == "__main__":
    run_integration_pipeline()
